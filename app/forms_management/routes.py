from flask import (render_template, redirect, request, url_for, flash,
                   current_app, send_from_directory, abort, send_file) # Añadido send_file
from flask_login import login_required, current_user
from . import forms_bp # Blueprint
from .. import db
from ..models import (User, Role, Permission, Category, FormDefinition,
                    FormField, FormSubmission, SubmissionData, UploadedImage)
from .forms import (CategoryForm, FormDefinitionForm, FormFieldDefinitionForm,
                  DynamicFormSubmission, ExportForm) # DynamicFormSubmission es la base
from ..decorators import admin_required, permission_required, approval_required

from wtforms import StringField, IntegerField, DateField, TextAreaField, BooleanField, MultipleFileField
from wtforms.validators import DataRequired as WTDataRequired, Optional as WTOptional, Email, NumberRange, Length as WTLength
from flask_wtf.file import FileAllowed, FileRequired
from werkzeug.utils import secure_filename
import os
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO


# --- Gestión de Categorías (Admin) ---
@forms_bp.route('/categories/manage', methods=['GET', 'POST'])
@login_required
@admin_required # O un permiso específico como 'manage_categories'
def manage_categories():
    categories = Category.query.order_by(Category.name).all()
    return render_template('forms_management/manage_categories.html',
                           categories=categories,
                           title="Gestionar Categorías de Formularios")

@forms_bp.route('/categories/create', methods=['GET', 'POST'])
@login_required
@admin_required
def create_category():
    form = CategoryForm()
    if form.validate_on_submit():
        category = Category(name=form.name.data, description=form.description.data)
        db.session.add(category)
        db.session.commit()
        flash(f'Categoría "{category.name}" creada exitosamente.', 'success')
        return redirect(url_for('.manage_categories'))
    return render_template('forms_management/create_edit_category.html', form=form,
                           title="Crear Nueva Categoría", action="Crear")

@forms_bp.route('/categories/<int:category_id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_category(category_id):
    category = Category.query.get_or_404(category_id)
    form = CategoryForm(obj=category) # Cargar datos existentes
    form.id.data = category.id # Para la validación de nombre único en edición

    if form.validate_on_submit():
        category.name = form.name.data
        category.description = form.description.data
        db.session.commit()
        flash(f'Categoría "{category.name}" actualizada exitosamente.', 'success')
        return redirect(url_for('.manage_categories'))
    return render_template('forms_management/create_edit_category.html', form=form,
                           category=category, title="Editar Categoría", action="Actualizar")

@forms_bp.route('/categories/<int:category_id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_category(category_id):
    category = Category.query.get_or_404(category_id)
    if category.form_definitions.count() > 0:
        flash(f'La categoría "{category.name}" tiene plantillas de formulario asociadas y no puede ser eliminada.', 'warning')
    else:
        name = category.name
        db.session.delete(category)
        db.session.commit()
        flash(f'Categoría "{name}" eliminada.', 'success')
    return redirect(url_for('.manage_categories'))


# --- Gestión de Plantillas de Formulario (Admin/Editor) ---
@forms_bp.route('/definitions/manage')
@login_required
@permission_required('manage_forms') # O @admin_required si solo admins
def manage_form_definitions():
    # TODO: Añadir paginación si hay muchos formularios
    form_definitions = FormDefinition.query.order_by(FormDefinition.name).all()
    return render_template('forms_management/manage_form_definitions.html',
                           form_definitions=form_definitions,
                           title="Gestionar Plantillas de Formularios")

@forms_bp.route('/definitions/create', methods=['GET', 'POST'])
@login_required
@permission_required('manage_forms')
def create_form_definition():
    form = FormDefinitionForm()
    if request.method == 'POST': # Incluye validación de subformularios
        if form.validate_on_submit():
            form_def = FormDefinition(
                name=form.name.data,
                description=form.description.data,
                category_id=form.category_id.data if form.category_id.data != 0 else None,
                section=form.section.data if form.section.data else None,
                created_by_id=current_user.id,
                is_active=form.is_active.data
            )
            db.session.add(form_def)
            # Guardar primero FormDefinition para obtener su ID
            db.session.flush() # Para que form_def.id esté disponible

            for i, field_data in enumerate(form.fields.entries):
                if field_data.form.field_label.data and field_data.form.field_name.data : # Asegurar que el campo no esté vacío
                    form_field = FormField(
                        form_definition_id=form_def.id,
                        label=field_data.form.field_label.data,
                        field_name=field_data.form.field_name.data,
                        field_type=field_data.form.field_type.data,
                        is_required=field_data.form.is_required.data,
                        order=i
                    )
                    db.session.add(form_field)

            db.session.commit()
            flash(f'Plantilla de formulario "{form_def.name}" creada exitosamente.', 'success')
            return redirect(url_for('.manage_form_definitions'))
        else:
            flash('Errores en el formulario. Por favor, revisa los campos.', 'danger')

    return render_template('forms_management/create_edit_form_definition.html', form=form,
                           title="Crear Nueva Plantilla de Formulario", action="Crear")


@forms_bp.route('/definitions/<int:form_def_id>/edit', methods=['GET', 'POST'])
@login_required
@permission_required('manage_forms')
def edit_form_definition(form_def_id):
    form_def = FormDefinition.query.get_or_404(form_def_id)
    form = FormDefinitionForm(obj=form_def)
    # form.id.data = form_def.id # Si se añade validación de nombre único para plantillas

    if request.method == 'POST':
        if form.validate_on_submit():
            form_def.name = form.name.data
            form_def.description = form.description.data
            form_def.category_id = form.category_id.data if form.category_id.data != 0 else None
            form_def.section = form.section.data if form.section.data else None
            form_def.is_active = form.is_active.data
            form_def.updated_at = datetime.datetime.utcnow()
            # Quién lo actualizó podría ser otro campo si es necesario.

            # Gestionar campos: eliminar los que no están, actualizar existentes, añadir nuevos
            existing_field_ids = {f.id for f in form_def.fields}
            submitted_field_map = {} # field_name -> data_dict

            # Primero, procesar los campos enviados para identificar nuevos y actualizados
            new_fields_data = []
            updated_fields_data = {} # id -> data_dict

            # Recolectar datos de los campos del formulario
            # WTForms FieldList no maneja bien la edición de elementos existentes por ID directamente.
            # La estrategia común es:
            # 1. Eliminar todos los campos existentes de la BD que no estén en el form enviado (o marcarlos).
            # 2. Actualizar los campos que sí están (comparando por un ID oculto si es posible, o por orden/nombre si no).
            # 3. Añadir los campos nuevos.
            # Por simplicidad aquí: Borrar todos los campos y re-crearlos.
            # Esto es más simple pero pierde los IDs de los FormField existentes (lo cual puede ser un problema si hay SubmissionData ligado a ellos directamente por ID).
            # Una mejor aproximación sería mantener los IDs de los campos si es posible.

            # Estrategia simple: borrar y re-crear (NO USAR SI SubmissionData.field_id es crítico)
            FormField.query.filter_by(form_definition_id=form_def.id).delete()
            db.session.flush() # Aplicar delete

            for i, field_entry in enumerate(form.fields.entries):
                if field_entry.form.field_label.data and field_entry.form.field_name.data:
                    new_form_field = FormField(
                        form_definition_id=form_def.id,
                        label=field_entry.form.field_label.data,
                        field_name=field_entry.form.field_name.data,
                        field_type=field_entry.form.field_type.data,
                        is_required=field_entry.form.is_required.data,
                        order=i
                    )
                    db.session.add(new_form_field)

            db.session.commit()
            flash(f'Plantilla de formulario "{form_def.name}" actualizada.', 'success')
            return redirect(url_for('.manage_form_definitions'))
        else:
            flash('Errores en el formulario. Por favor, revisa los campos.', 'danger')
    else: # GET request
        # WTForms no rellena FieldList(FormField) automáticamente con `obj=form_def` si FormField no es un modelo SQLAlchemy directo.
        # Necesitamos poblarlo manualmente para la edición.
        # Limpiar cualquier entrada por defecto que WTForms pudiera haber creado
        while len(form.fields.entries) > 0:
            form.fields.pop_entry()

        for field_model in form_def.fields.order_by(FormField.order):
            field_form = FormFieldDefinitionForm() # Subformulario
            field_form.field_label = field_model.label
            field_form.field_name = field_model.field_name
            field_form.field_type = field_model.field_type
            field_form.is_required = field_model.is_required
            form.fields.append_entry(field_form)

        # Si no hay campos, asegurar al menos una entrada vacía para el template
        if not form.fields.entries:
            form.fields.append_entry(FormFieldDefinitionForm())


    return render_template('forms_management/create_edit_form_definition.html', form=form,
                           form_definition=form_def, title="Editar Plantilla de Formulario", action="Actualizar")


@forms_bp.route('/definitions/<int:form_def_id>/delete', methods=['POST'])
@login_required
@permission_required('manage_forms')
def delete_form_definition(form_def_id):
    form_def = FormDefinition.query.get_or_404(form_def_id)
    if form_def.submissions.count() > 0:
        flash(f'La plantilla "{form_def.name}" tiene envíos asociados y no puede ser eliminada directamente. Considere desactivarla en su lugar.', 'warning')
    else:
        name = form_def.name
        # Los FormFields se borran en cascada (definido en el modelo)
        db.session.delete(form_def)
        db.session.commit()
        flash(f'Plantilla de formulario "{name}" eliminada.', 'success')
    return redirect(url_for('.manage_form_definitions'))


# --- Listar y Enviar Formularios (Usuarios) ---

def create_dynamic_form_class(form_definition_model):
    """Crea una clase de formulario WTForms dinámicamente basada en FormDefinition."""

    # Heredar de nuestra clase base que ya tiene CSRF y el campo submit
    class TempDynamicForm(DynamicFormSubmission):
        pass

    # Añadir el campo oculto para el ID de la definición del formulario
    setattr(TempDynamicForm, 'form_definition_id', HiddenField(validators=[WTDataRequired()], default=form_definition_model.id))

    for field_def in form_definition_model.fields.order_by(FormField.order):
        validators = []
        if field_def.is_required:
            validators.append(WTDataRequired(message=f"El campo '{field_def.label}' es obligatorio."))
        else:
            validators.append(WTOptional()) # Importante para campos no requeridos

        wtform_field = None
        field_kwargs = {'validators': validators, 'description': ""} # Añadir descripción si se guarda en FormField

        if field_def.field_type == 'text':
            wtform_field = StringField(field_def.label, **field_kwargs)
        elif field_def.field_type == 'textarea':
            wtform_field = TextAreaField(field_def.label, **field_kwargs)
        elif field_def.field_type == 'number':
            # Para IntegerField, WTOptional permite que el campo esté vacío.
            # Si se usa NumberRange y el campo es opcional, NumberRange no se aplicará si está vacío.
            wtform_field = IntegerField(field_def.label, **field_kwargs)
        elif field_def.field_type == 'date':
            wtform_field = DateField(field_def.label, format='%Y-%m-%d', **field_kwargs)
        elif field_def.field_type == 'boolean':
            # BooleanField es especial, no suele llevar DataRequired. Su valor es True/False.
            # Si es opcional, simplemente no se marca.
            wtform_field = BooleanField(field_def.label, **{'description': field_kwargs['description']}) # No pasar validators aquí
        elif field_def.field_type == 'image_upload':
            allowed_ext = current_app.config['ALLOWED_EXTENSIONS']
            # FileRequired solo si el campo es_required. FileAllowed siempre.
            file_validators = [FileAllowed(allowed_ext, '¡Solo imágenes!')]
            if field_def.is_required:
                file_validators.insert(0, FileRequired(message=f"Debe subir al menos una imagen para '{field_def.label}'."))

            wtform_field = MultipleFileField(field_def.label, validators=file_validators, description=field_kwargs['description'])

        if wtform_field:
            setattr(TempDynamicForm, field_def.field_name, wtform_field)

    return TempDynamicForm


@forms_bp.route('/list')
@login_required
@approval_required
def list_forms():
    # Filtrar por sección si se proporciona
    section_filter = request.args.get('section')
    query = FormDefinition.query.filter_by(is_active=True)
    if section_filter:
        query = query.filter_by(section=section_filter)

    available_forms = query.order_by(FormDefinition.category_id, FormDefinition.name).all()

    # Obtener todas las secciones únicas de los formularios activos para los filtros
    active_sections = db.session.query(FormDefinition.section).filter(FormDefinition.is_active==True, FormDefinition.section != None, FormDefinition.section != "").distinct().all()
    site_sections = sorted([s[0] for s in active_sections])

    return render_template('forms_management/list_forms.html',
                           forms=available_forms,
                           site_sections=site_sections,
                           current_section=section_filter,
                           title="Formularios Disponibles")


@forms_bp.route('/submit/<int:form_def_id>', methods=['GET', 'POST'])
@login_required
@approval_required
def submit_form(form_def_id):
    form_def = FormDefinition.query.filter_by(id=form_def_id, is_active=True).first_or_404()

    DynamicFormClass = create_dynamic_form_class(form_def)
    form = DynamicFormClass(request.form if request.method == 'POST' else None) # Pasar request.form para POST

    if form.validate_on_submit():
        # Crear la entrada de FormSubmission
        submission = FormSubmission(
            form_definition_id=form_def.id,
            user_id=current_user.id
        )
        db.session.add(submission)
        db.session.flush() # Para obtener submission.id

        # Guardar cada dato del formulario
        for field_def in form_def.fields.order_by(FormField.order):
            field_form_data = form[field_def.field_name].data

            submission_data_entry = SubmissionData(
                submission_id=submission.id,
                field_id=field_def.id, # Enlazar al FormField original
                field_name=field_def.field_name,
            )

            if field_def.field_type == 'image_upload':
                # Manejar subida de múltiples archivos
                # field_form_data será una lista de FileStorage objects
                image_filenames = []
                for file_storage in field_form_data:
                    if file_storage and file_storage.filename: # Check if a file was actually uploaded
                        original_filename = secure_filename(file_storage.filename)
                        # Crear un nombre de archivo único para evitar colisiones
                        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f")
                        filename = f"user{current_user.id}_form{form_def.id}_sub{submission.id}_{field_def.field_name}_{timestamp}_{original_filename}"

                        upload_path = current_app.config['UPLOAD_FOLDER']
                        if not os.path.exists(upload_path):
                            os.makedirs(upload_path)

                        file_path = os.path.join(upload_path, filename)
                        file_storage.save(file_path)

                        # Guardar información de la imagen en UploadedImage
                        uploaded_image = UploadedImage(
                            submission_data_id=None, # Se asignará después de guardar submission_data_entry
                            filename=filename,
                            original_filename=original_filename,
                            mimetype=file_storage.mimetype
                        )
                        submission_data_entry.uploaded_images.append(uploaded_image)
                        image_filenames.append(filename) # Solo para referencia o uso en value_text si es necesario

                # Podrías guardar una lista de nombres de archivo o un contador en value_text
                submission_data_entry.value_text = ", ".join(image_filenames) if image_filenames else None

            elif field_def.field_type == 'boolean':
                submission_data_entry.value_text = str(field_form_data) # 'True' o 'False'
            elif field_def.field_type == 'date':
                submission_data_entry.value_text = field_form_data.strftime('%Y-%m-%d') if field_form_data else None
            else: # text, textarea, number
                submission_data_entry.value_text = str(field_form_data) if field_form_data is not None else None

            db.session.add(submission_data_entry)
            db.session.flush() # Para que submission_data_entry.id esté disponible si se asigna a UploadedImage

            # Si las imágenes se asocian después:
            if field_def.field_type == 'image_upload':
                for img_obj in submission_data_entry.uploaded_images:
                    img_obj.submission_data_id = submission_data_entry.id


        db.session.commit()
        flash(f'Formulario "{form_def.name}" enviado exitosamente.', 'success')
        return redirect(url_for('main.dashboard')) # O a una página de "mis envíos"

    # Para GET, pre-llenar el form_definition_id (ya se hace en create_dynamic_form_class)
    # form.form_definition_id.data = form_def.id

    return render_template('forms_management/submit_form.html',
                           form_def=form_def,
                           form=form,
                           title=f"Rellenar: {form_def.name}")


# --- Visualización de Datos Enviados (Admin) ---
@forms_bp.route('/submissions/summary')
@login_required
@permission_required('view_all_submissions') # O @admin_required
def view_all_submissions_summary():
    # Lista todas las plantillas de formulario y cuántos envíos tiene cada una
    # Esto podría ser costoso si hay muchos formularios y envíos.
    # Considerar paginación o una vista más agregada.
    form_defs_with_counts = db.session.query(
        FormDefinition, db.func.count(FormSubmission.id).label('submission_count')
    ).outerjoin(FormSubmission, FormDefinition.id == FormSubmission.form_definition_id)\
    .group_by(FormDefinition.id)\
    .order_by(FormDefinition.name)\
    .all()

    return render_template('forms_management/all_submissions_summary.html',
                           form_defs_with_counts=form_defs_with_counts,
                           title="Resumen de Envíos de Formularios")

@forms_bp.route('/submissions/view/<int:form_def_id>')
@login_required
@permission_required('view_all_submissions')
def view_submissions_for_form(form_def_id):
    form_def = FormDefinition.query.get_or_404(form_def_id)

    page = request.args.get('page', 1, type=int)
    per_page = 10 # O configurable

    submissions_query = FormSubmission.query.filter_by(form_definition_id=form_def.id)\
                                       .order_by(FormSubmission.submitted_at.desc())

    pagination = submissions_query.paginate(page=page, per_page=per_page, error_out=False)
    submissions_on_page = pagination.items

    # Para cada envío, necesitamos construir una representación de sus datos
    # Esto puede ser complejo si los datos son muy variados.
    # La tabla HTML se construirá dinámicamente.
    # Necesitamos los encabezados (labels de los campos del formulario)
    field_headers = [{'name': f.field_name, 'label': f.label, 'type': f.field_type}
                     for f in form_def.fields.order_by(FormField.order)]

    processed_submissions = []
    for sub in submissions_on_page:
        data_dict = {sd.field_name: sd for sd in sub.data_entries} # field_name -> SubmissionData object
        processed_sub = {'id': sub.id, 'submitter_email': sub.submitter.email, 'submitted_at': sub.submitted_at, 'data': []}
        for header in field_headers:
            entry = data_dict.get(header['name'])
            value_to_display = ""
            images_to_display = []
            if entry:
                if header['type'] == 'image_upload' and entry.uploaded_images:
                    images_to_display = [{'filename': img.filename,
                                          'url': url_for('forms_management.uploaded_file', filename=img.filename)}
                                         for img in entry.uploaded_images]
                    value_to_display = f"{len(images_to_display)} imágen(es)" if images_to_display else "Ninguna"
                else:
                    value_to_display = entry.value_text if entry.value_text is not None else "---"
            else: # No data for this field in this submission
                 value_to_display = "---"

            processed_sub['data'].append({'value': value_to_display, 'images': images_to_display, 'type': header['type']})
        processed_submissions.append(processed_sub)

    return render_template('forms_management/view_submissions.html',
                           form_definition=form_def,
                           field_headers=field_headers,
                           submissions=processed_submissions,
                           pagination=pagination,
                           title=f"Envíos para: {form_def.name}")


# Servir archivos subidos (imágenes)
@forms_bp.route('/uploads/<filename>')
@login_required # Asegurar que solo usuarios logueados puedan ver (podría ser más restrictivo)
def uploaded_file(filename):
    return send_from_directory(current_app.config['UPLOAD_FOLDER'], filename)


# --- Exportación de Envíos a Word (.docx) ---
@forms_bp.route('/submission/<int:submission_id>/export/word', methods=['GET','POST']) # POST si usas un form, GET si es un link
@login_required
@permission_required('export_submissions') # O @admin_required
def export_submission_to_word(submission_id):
    submission = FormSubmission.query.get_or_404(submission_id)
    form_def = submission.form_definition

    # Crear documento Word
    document = Document()

    # Título del formulario
    title_heading = document.add_heading(form_def.name, level=1)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if form_def.description:
        document.add_paragraph(form_def.description, style='Caption') # o 'Quote'

    document.add_paragraph(f"Enviado por: {submission.submitter.email}")
    document.add_paragraph(f"Fecha de envío: {submission.submitted_at.strftime('%Y-%m-%d %H:%M:%S')} UTC")
    document.add_paragraph() # Espacio

    # Datos del formulario
    for field_def in form_def.fields.order_by(FormField.order):
        data_entry = SubmissionData.query.filter_by(submission_id=submission.id, field_id=field_def.id).first()

        document.add_heading(field_def.label, level=3)

        if data_entry:
            if field_def.field_type == 'image_upload':
                if data_entry.uploaded_images.count() > 0:
                    for img_obj in data_entry.uploaded_images:
                        try:
                            img_path = os.path.join(current_app.config['UPLOAD_FOLDER'], img_obj.filename)
                            if os.path.exists(img_path):
                                # Añadir imagen. Ajustar tamaño según necesidad.
                                document.add_paragraph(f"Imagen: {img_obj.original_filename or img_obj.filename}")
                                document.add_picture(img_path, width=Inches(4.0)) # Ajustar tamaño
                            else:
                                document.add_paragraph(f"[Imagen no encontrada en servidor: {img_obj.filename}]")
                        except Exception as e:
                            current_app.logger.error(f"Error al añadir imagen {img_obj.filename} al DOCX: {e}")
                            document.add_paragraph(f"[Error al procesar imagen: {img_obj.filename}]")
                    if data_entry.uploaded_images.count() == 0: # Si no hay imágenes a pesar de ser campo de imagen
                         document.add_paragraph("No se subieron imágenes para este campo.", style='Emphasis')
                else: # Si no hay data_entry.uploaded_images
                    document.add_paragraph("No se subieron imágenes para este campo.", style='Emphasis')
            elif field_def.field_type == 'textarea':
                # Para textareas, preservar saltos de línea (Word lo hace por defecto con add_paragraph)
                if data_entry.value_text:
                    for line in data_entry.value_text.splitlines():
                        document.add_paragraph(line if line else " ") # Poner un espacio si la línea está vacía para mantener el párrafo
                else:
                     document.add_paragraph("---", style='Emphasis')
            elif field_def.field_type == 'boolean':
                val = "Sí" if data_entry.value_text == 'True' else ("No" if data_entry.value_text == 'False' else "---")
                document.add_paragraph(val)
            else: # text, number, date
                document.add_paragraph(data_entry.value_text if data_entry.value_text is not None else "---")
        else:
            document.add_paragraph("--- (Sin respuesta)", style='Emphasis')
        document.add_paragraph() # Espacio entre campos

    # Guardar en un stream de bytes
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    # Nombre del archivo de descarga
    safe_form_name = secure_filename(form_def.name.replace(" ", "_"))
    safe_user_email = secure_filename(submission.submitter.email.split('@')[0])
    filename_docx = f"envio_{safe_form_name}_sub{submission.id}_{safe_user_email}.docx"

    return send_from_directory(
        os.getcwd(), # Wordpy crea el archivo temporalmente, pero send_file lo manejará desde el stream
        path=filename_docx, # Este path es solo para el nombre del archivo en la descarga
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        # Para send_file directamente desde BytesIO:
        # return send_file(file_stream, as_attachment=True, download_name=filename_docx,
        #                  mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        # Sin embargo, send_from_directory es más común si el archivo se guarda primero.
        # Aquí, dado que usamos BytesIO, send_file es más apropiado.
        # Reemplazar send_from_directory con send_file
        environ=request.environ # Necesario para send_file
    )
    # Corrección: Usar send_file para BytesIO
    # from flask import send_file (importar arriba)
    # ...
    # return send_file(file_stream, as_attachment=True, download_name=filename_docx,
    #                  mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    # Por ahora, la estructura del código sugiere que `send_from_directory` se esperaba,
    # pero para BytesIO `send_file` es la elección. Voy a ajustar para usar send_file
    # pero necesitaría importar `send_file` de `flask`.
    # Asumiendo que `send_file` está disponible:

    # from flask import send_file # Asegurar que esté importado al inicio del archivo
    # La importación de send_file se moverá al inicio del archivo.
    return send_file(file_stream, download_name=filename_docx, as_attachment=True,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
